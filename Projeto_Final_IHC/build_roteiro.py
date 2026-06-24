# -*- coding: utf-8 -*-
"""Roteiro de apresentação (fala por fala) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE=os.path.dirname(os.path.abspath(__file__))
INK=RGBColor(0x1A,0x17,0x14); CLAY=RGBColor(0xA8,0x4E,0x30); SAGE=RGBColor(0x5F,0x6B,0x50); GREY=RGBColor(0x6b,0x63,0x59)
doc=Document(); st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11); st.font.color.rgb=INK
for s in doc.sections: s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)

def slide(n,title,who,tempo,fala,acao=None):
    # cabeçalho do slide em tabela 1 célula
    tb=doc.add_table(rows=1,cols=1); tb.style="Table Grid"; c=tb.rows[0].cells[0]; shade(c,"1A1714")
    pr=c.paragraphs[0]
    r=pr.add_run(f"SLIDE {n} · {title}"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xF6,0xF1,0xE9)
    pr2=c.add_paragraph();
    whoCol = CLAY if "Arthur" in who else (SAGE if "Sávio" in who else RGBColor(0xE7,0xC4,0xB2))
    r2=pr2.add_run(f"{who}   ·   ⏱ {tempo}"); r2.bold=True; r2.font.size=Pt(10); r2.font.color.rgb=RGBColor(0xE7,0xC4,0xB2)
    # fala
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    rf=p.add_run("Fala: "); rf.bold=True; rf.font.color.rgb=CLAY
    p.add_run(fala)
    if acao:
        pa=doc.add_paragraph(); pa.paragraph_format.space_after=Pt(10)
        ra=pa.add_run("► Ação na tela: "); ra.bold=True; ra.font.color.rgb=SAGE
        ra2=pa.add_run(acao); ra2.italic=True
    else:
        doc.add_paragraph().paragraph_format.space_after=Pt(6)

# capa do roteiro
tt=doc.add_paragraph(); tt.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tt.add_run("Roteiro de Apresentação — Stylo"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=INK
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=sub.add_run("Avaliação de Usabilidade · IHC · Arthur Faria Estrela & Sávio Issa de Sousa · ~12 min30"); rs.italic=True; rs.font.size=Pt(11); rs.font.color.rgb=CLAY
note=doc.add_paragraph(); note.alignment=WD_ALIGN_PARAGRAPH.CENTER
rn=note.add_run("Dica: não leia o slide nem este papel palavra por palavra — use como apoio. Olhe para a plateia, "
                "fale natural. Os trechos entre [colchetes] são lembretes, não para falar."); rn.font.size=Pt(9.5); rn.font.color.rgb=GREY
doc.add_paragraph()

# ============================ FALAS ============================
slide(1,"Capa","Arthur","0:30",
 "Boa noite, professora e colegas. Nós somos o Arthur e o Sávio. Hoje vamos apresentar a avaliação de "
 "usabilidade do Stylo — uma plataforma de gestão e agendamento para o setor de estética que nós "
 "mesmos desenvolvemos. A proposta foi pegar o nosso próprio sistema e colocá-lo à prova com usuários "
 "reais, aplicando as técnicas de IHC que vimos na disciplina. No final, transformamos os problemas "
 "encontrados em um protótipo de melhoria, que vamos mostrar funcionando ao vivo.")

slide(2,"O produto","Arthur","1:00",
 "Primeiro, o que é o Stylo. É um SaaS — um software como serviço — voltado para barbearias, salões e "
 "clínicas de estética que ainda usam agenda de papel ou anotações no WhatsApp. Ele tem dois lados: "
 "uma página pública, onde o cliente final agenda o horário sozinho pelo celular; e um painel de "
 "gestão, onde o dono organiza a equipe, os serviços, a disponibilidade e acompanha o faturamento. O "
 "grande destaque é o agendamento inteligente, que bloqueia automaticamente horários em conflito. "
 "Como são dois perfis bem diferentes de usuário — o profissional e o cliente —, avaliar a usabilidade "
 "ficou ainda mais importante.")

slide(3,"Framework DECIDE","Arthur","0:50",
 "Para não avaliar no 'achismo', a gente estruturou tudo com o framework DECIDE, do Preece, Rogers e "
 "Sharp. Cada letra é uma etapa: determinamos os objetivos — facilidade de uso, eficiência e "
 "satisfação; exploramos as perguntas — será que o usuário conclui as tarefas sozinho?; escolhemos os "
 "métodos; identificamos as questões práticas, como o perfil dos participantes; decidimos as questões "
 "éticas, com termo de consentimento e anonimato; e por fim avaliamos os dados. Esse framework foi o "
 "fio condutor de todo o trabalho.")

slide(4,"Métodos","Arthur","0:40",
 "Os métodos que escolhemos se complementam. O 'Pensar em voz alta', em que o usuário narra o "
 "raciocínio enquanto usa o sistema — isso revela a confusão na hora que ela acontece. O SUS, que é um "
 "questionário padronizado que vira uma nota de zero a cem. E uma entrevista no final, com perguntas "
 "abertas. Ou seja: a gente mediu a usabilidade com número, e entendeu o porquê com a fala do usuário.")

slide(5,"Avaliação heurística","Arthur","0:50",
 "Antes de chamar os usuários, fizemos uma avaliação heurística, usando as dez heurísticas do Nielsen "
 "para inspecionar as telas. Encontramos quatorze problemas. Os mais sérios: a tela de login não tinha "
 "'esqueci minha senha', o que prende o usuário; o calendário não mostrava em quais dias havia vaga, "
 "obrigando a clicar dia por dia; e botões que ficavam ativos mesmo com campo inválido. Esse gráfico à "
 "direita mostra os problemas separados por gravidade.")

slide(6,"Experimento","Arthur","0:50",
 "Aí partimos para o teste com gente de verdade. Recrutamos seis participantes do setor de estética, do "
 "iniciante ao avançado — donos de salão, barbeiros, gerente de clínica. O formato foi híbrido: "
 "presencial e remoto com gravação de tela. Eles executaram oito tarefas no painel. Antes disso, "
 "fizemos um teste piloto com uma pessoa, que durou vinte e três minutos e já nos avisou de dois "
 "ajustes: a dificuldade de achar o menu da equipe e a falta de dados nos gráficos — que a gente "
 "corrigiu antes do teste real.")

slide(7,"Resultados — tarefas","Arthur","0:50",
 "E os resultados. Esses dois gráficos resumem o desempenho nas oito tarefas. A boa notícia: nenhuma "
 "tarefa ficou incompleta — todo mundo conseguiu fazer tudo. Mas dois pontos chamaram atenção, em "
 "laranja: a Tarefa 2, adicionar um profissional à equipe, e a Tarefa 4, definir a disponibilidade. "
 "Foram as que mais precisaram de ajuda e as que levaram mais tempo. Já guardem esse 'adicionar "
 "profissional', porque ele volta na proposta de melhoria.")

slide(8,"Resultados — SUS","Arthur","0:40",
 "E a satisfação geral, medida pelo SUS, deu setenta e oito vírgula três. Para terem ideia, a média de "
 "mercado é sessenta e oito, e acima de oitenta já é considerado 'excelente'. Ou seja, o Stylo ficou "
 "na faixa 'bom', quase 'excelente'. É uma base sólida — os problemas que achamos são de refinamento, "
 "não de algo quebrado. E agora o Sávio assume para falar do lado qualitativo. [Passa para o Sávio]")

slide(9,"Think-Aloud","Sávio","0:50",
 "Obrigado, Arthur. Esse número bom esconde detalhes que só a fala revela. Na sessão de 'pensar em voz "
 "alta' no agendamento, o usuário disse coisas como: 'o preço tá com a fonte pequena, difícil de ler'; "
 "'fiquei na dúvida se o horário estava bloqueado'; e o mais marcante: 'cliquei em confirmar e fui "
 "jogado pra tela de login sem aviso nenhum'. Isso quebrou a experiência. Mas teve um ponto muito "
 "positivo: mesmo nesse susto, o sistema salvou todas as escolhas dele — não perdeu nada. Isso evitou "
 "um problema crítico.")

slide(10,"Discussão","Sávio","0:50",
 "Juntando tudo — heurística, fala e teste — chegamos a um diagnóstico: o Stylo é forte no núcleo e tem "
 "atrito nas bordas. De um lado, pontos fortes claros: o agendamento inteligente foi elogiado por "
 "todos, e a preservação dos dados mostrou maturidade. Do outro, quatro problemas prioritários: a "
 "equipe escondida nas configurações, o redirecionamento sem aviso, o calendário sem indicador de "
 "vaga e o contraste fraco dos preços. Foram esses quatro que viraram o nosso plano de redesenho.")

slide(11,"Proposta de melhorias","Sávio","0:50",
 "Para cada problema, uma solução concreta. A equipe vira um item de primeiro nível no menu, com botão "
 "de adicionar em destaque. O preço ganha fonte maior e cor escura. O redirecionamento passa a ter um "
 "aviso antes, garantindo que os dados estão salvos. O calendário ganha pontos coloridos de "
 "disponibilidade e um atalho de 'próximo horário livre'. E voltamos com o 'esqueci minha senha'. Mas "
 "em vez de só falar, a gente construiu isso — deixa eu mostrar funcionando.")

slide(12,"Protótipo — Equipe","Sávio","DEMO 1:15",
 "[Abrir prototipo/index.html] Esse é o protótipo navegável, com as telas reais redesenhadas. Repara "
 "que aqui na barra lateral a 'Equipe' agora é um item de primeiro nível — não está mais escondida na "
 "engrenagem. E o botão 'adicionar profissional' fica em destaque, bem no topo. Lembram que essa era a "
 "tarefa que mais travou? Resolvida. [Clicar em 'Mostrar melhorias' para destacar] Esse botão aqui em "
 "cima destaca cada melhoria que fizemos sobre a interface.",
 "Abrir o protótipo no navegador. Clicar em 'Painel · Equipe'. Mostrar a sidebar e o botão. Ligar 'Mostrar melhorias'.")

slide(13,"Protótipo — Agendamento e calendário","Sávio","DEMO 1:15",
 "Agora o lado do cliente. [Ir em 'Agendamento'] Olhem o preço: fonte grande, alto contraste, fácil de "
 "ler — era a reclamação do nosso usuário. Sigo o fluxo... e aqui o calendário: cada dia tem uma "
 "bolinha mostrando se tem bastante vaga, pouca, ou nenhuma — sem precisar clicar um por um. E esse "
 "atalho, 'próximo horário livre', pula direto pra primeira vaga. Por fim, quando confirmo o "
 "agendamento... [clicar em confirmar] aparece esse aviso, explicando que precisa criar conta e "
 "garantindo que nada foi perdido. Nada de susto.",
 "Ir em 'Agendamento (cliente)'. Mostrar preços. Avançar até o calendário (passo 3). Apontar bolinhas e 'próximo horário livre'. Clicar em 'Confirmar agendamento' para abrir o modal de aviso.")

slide(14,"Validação do protótipo","Sávio","0:40",
 "E não ficou só no 'achamos bonito'. Fizemos um teste rápido do protótipo com três usuários, "
 "comparando antes e depois. Achar a 'equipe' caiu de cento e dez segundos para vinte e dois. Achar um "
 "dia com vaga no calendário, de quarenta para oito segundos. E ninguém mais foi pego de surpresa pelo "
 "login. Ou seja, as mudanças funcionaram de verdade.")

slide(15,"Conclusão","Sávio","0:40",
 "Concluindo: o Stylo já nasceu bom, com SUS de setenta e oito e nenhuma tarefa não concluída. Com seis "
 "melhorias simples e baratas de implementar — mexer no menu, contraste, avisos e indicadores visuais "
 "— dá pra levar ele de 'bom' para 'excelente'. E o mais importante: isso só ficou claro porque a "
 "gente ouviu o usuário, em vez de adivinhar. Esse é o valor do design centrado no usuário.")

slide(16,"Encerramento","Arthur & Sávio","0:30",
 "É isso! O Stylo, avaliado, diagnosticado e com um caminho claro de evolução. Agradecemos a atenção e "
 "estamos à disposição para perguntas.")

out=os.path.join(HERE,"Roteiro_Apresentacao.docx")
doc.save(out); print("salvo:",out)
