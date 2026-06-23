# -*- coding: utf-8 -*-
"""
Gera o Projeto Final de IHC (Avaliação da Usabilidade do Stylo) em .docx e .html
a partir de uma única fonte de conteúdo.
"""
import os, html as _html

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")

# =====================================================================
#  CABEÇALHO / METADADOS
# =====================================================================
META = {
    "instituicao": "INSTITUTO FEDERAL GOIANO – CAMPUS URUTAÍ",
    "curso": "Sistemas de Informação",
    "turma": "7º Período",
    "disciplina": "Interface Humano-Computador (IHC)",
    "professora": "Cristiane de Fátima dos Santos Cardoso",
    "alunos": "Arthur Faria Estrela e Sávio Issa de Sousa",
    "data": "16 de junho de 2026",
    "titulo": "Avaliação da Usabilidade do Stylo",
    "subtitulo": "Plataforma SaaS de gestão e agendamento para o setor de estética",
}

# =====================================================================
#  CONTEÚDO  (lista de blocos)
#  tipos: h1, h2, p, bul, num, tbl, img, quote, callout, pagebreak
#  Em p/bul/num/quote/callout use **negrito** para destaque.
# =====================================================================
C = []
def h1(t): C.append(("h1", t))
def h2(t): C.append(("h2", t))
def p(t):  C.append(("p", t))
def bul(items): C.append(("bul", items))
def num(items): C.append(("num", items))
def tbl(headers, rows): C.append(("tbl", headers, rows))
def img(rel, cap): C.append(("img", rel, cap))
def quote(t, who): C.append(("quote", t, who))
def callout(title, t): C.append(("callout", title, t))
def pb(): C.append(("pagebreak",))

# ----------------------------------------------------------- RESUMO
p("**Resumo.** Este trabalho apresenta a avaliação de usabilidade da plataforma "
  "**Stylo**, um SaaS de gestão e agendamento para o setor de estética. A avaliação "
  "foi estruturada com o framework **DECIDE** (Preece, Rogers e Sharp) e combinou três "
  "técnicas complementares: **avaliação heurística** (10 heurísticas de Nielsen), o "
  "protocolo **“Pensar em voz alta” (Think-Aloud)** e um **teste de usabilidade** com 6 "
  "participantes do público-alvo, encerrado com o questionário padronizado **SUS (System "
  "Usability Scale)**. A plataforma obteve **SUS = 78,3** (classificação “Bom”, próximo ao "
  "limiar de “Excelente”), com fluxo principal eficiente e o algoritmo de agendamento "
  "inteligente como ponto forte unânime. Os principais atritos — arquitetura da informação "
  "para a gestão de equipe, contraste de preços, redirecionamento sem aviso para login e "
  "ausência de indicadores de disponibilidade no calendário — foram traduzidos em um "
  "**protótipo navegável de alta fidelidade**, submetido a um breve teste de validação.")

# =====================================================================
# 1. PRODUTO A SER AVALIADO
# =====================================================================
h1("1. Produto a ser avaliado")
p("O **Stylo** é uma plataforma **SaaS (Software as a Service)** de **gestão e agendamento** "
  "voltada ao setor de estética e beleza. O sistema atende estabelecimentos como "
  "**barbearias, salões de beleza e clínicas de estética** que desejam substituir métodos "
  "manuais (agenda de papel, anotações em aplicativos de mensagem) por uma solução digital "
  "centralizada de marcação de horários, organização da equipe e acompanhamento financeiro.")
p("Tecnicamente, o produto é uma **aplicação web responsiva**, com front-end em "
  "**React + TypeScript (Vite)** e back-end em **Java/Spring**, integrando ainda notificações "
  "via **WhatsApp**. A interface adapta-se a desktops, notebooks e smartphones, cobrindo tanto "
  "o uso no balcão do estabelecimento quanto o agendamento feito pelo cliente em seu próprio "
  "celular.")

h2("1.1 Principais características e funcionalidades")
bul([
  "**Página pública de agendamento:** fluxo em que o cliente final escolhe o serviço, o "
  "profissional, a data e o horário e confirma a marcação.",
  "**Agendamento inteligente (anti-conflito):** algoritmo que bloqueia automaticamente horários "
  "já ocupados, evitando sobreposição de atendimentos.",
  "**Painel de gestão:** visão geral de agendamentos, cadastro de clientes, gestão da equipe, "
  "cadastro de serviços e definição de disponibilidade (dias e horários).",
  "**Métricas de faturamento:** gráficos de faturamento projetado por período (dia/semana/mês).",
  "**Gestão de status do atendimento:** acompanhamento do ciclo de vida do agendamento "
  "(aguardando aceite → confirmado → concluído).",
])

h2("1.2 Usuários típicos")
p("A plataforma possui **dois perfis de usuário** claramente distintos, o que torna a avaliação "
  "de usabilidade especialmente relevante:")
bul([
  "**Profissionais e gestores de estética** (donos de barbearia, salões e clínicas, gerentes e "
  "atendentes): utilizam o **painel de gestão** para administrar agenda, equipe, serviços e "
  "finanças. Possuem níveis variados de familiaridade com tecnologia.",
  "**Clientes finais:** utilizam a **página pública** para agendar um serviço de forma rápida, "
  "muitas vezes pela primeira vez e a partir do celular.",
])

h2("1.3 Contextos de uso")
bul([
  "**No estabelecimento, durante a rotina:** o profissional consulta e gerencia a agenda entre um "
  "atendimento e outro, frequentemente com pouco tempo e em telas menores.",
  "**Em mobilidade:** donos de rede acompanham faturamento e solicitações pelo smartphone, fora "
  "do local de trabalho.",
  "**Cliente agendando à distância:** marcação feita de casa ou na rua, sem treinamento prévio e "
  "sob a expectativa de uma experiência tão simples quanto a de um aplicativo de delivery.",
])

# =====================================================================
# 2. FRAMEWORK DECIDE
# =====================================================================
pb()
h1("2. Framework DECIDE")
p("O **DECIDE** é um framework proposto por **Preece, Rogers e Sharp (2005)** para planejar e "
  "conduzir avaliações de usabilidade de forma sistemática. O acrônimo organiza a avaliação em "
  "seis etapas, garantindo que objetivos, métodos, questões práticas e éticas sejam definidos "
  "antes da coleta de dados. As etapas são: **D**eterminar os objetivos; **E**xplorar as "
  "perguntas; escolher (**C**hoose) os métodos; **I**dentificar questões práticas; **D**ecidir "
  "sobre questões éticas; e **E**valiar, interpretar e apresentar os dados.")
p("A seguir, apresenta-se o que foi obtido em cada etapa para o **Stylo**.")

tbl(["Etapa", "Aplicação ao Stylo"],
[
 ["D — Determinar os objetivos",
  "Avaliar a facilidade de uso, a eficiência na execução das tarefas e a satisfação geral do "
  "usuário ao utilizar as funções principais da plataforma (agendamento público e gestão)."],
 ["E — Explorar as perguntas",
  "Os usuários conseguem concluir o fluxo principal de forma independente? A navegação do painel "
  "é intuitiva? A terminologia é compreendida? O cliente consegue agendar sem treinamento?"],
 ["C — Escolher os métodos",
  "Think-Aloud (desempenho real e modelos mentais), SUS (medida quantitativa padronizada) e "
  "entrevista semiestruturada (preferências, reações e frustrações)."],
 ["I — Identificar questões práticas",
  "Perfil: profissionais/gestores de estética e clientes finais, com níveis variados de "
  "experiência. Recrutamento de 6 participantes (faixa de 5 a 12 recomendada). Ambiente híbrido "
  "(presencial e remoto), em ambiente de homologação controlado."],
 ["D — Decidir sobre ética",
  "Termo de consentimento e autorização de gravação; garantia de anonimato nos relatórios; "
  "confidencialidade dos dados de áudio/vídeo; e direito de interromper a avaliação a qualquer "
  "momento."],
 ["E — Avaliar e interpretar",
  "Codificação dos incidentes do Think-Aloud, tabulação da matriz de tarefas, cálculo do SUS e "
  "triangulação com a avaliação heurística (Seções 5 e 6)."],
])

h2("2.1 Justificativa dos métodos escolhidos")
bul([
  "**Pensar em voz alta (Think-Aloud):** método mais direto para coletar dados sobre o desempenho "
  "real nas tarefas, revelando problemas de navegação, modelos mentais incorretos e confusões na "
  "interface enquanto o usuário verbaliza seu raciocínio.",
  "**Questionário SUS:** instrumento padronizado, quantitativo e academicamente reconhecido. São "
  "10 afirmações em escala Likert de 1 a 5, convertidas em uma pontuação final de 0 a 100, "
  "permitindo comparar a usabilidade com benchmarks consolidados.",
  "**Entrevista semiestruturada (pós-teste):** complementa os dados numéricos com perguntas "
  "abertas (ex.: “O que você achou da organização do menu?”), evitando induzir respostas e "
  "capturando preferências e frustrações.",
])

# =====================================================================
# 3. AVALIAÇÃO HEURÍSTICA
# =====================================================================
pb()
h1("3. Avaliação heurística")
p("A avaliação heurística foi conduzida com base nas **10 heurísticas de usabilidade de Jakob "
  "Nielsen**. Dois avaliadores percorreram as telas críticas do Stylo (login, página pública de "
  "agendamento, seleção de data/hora, listagem de profissionais e painel) registrando violações. "
  "O quadro a seguir consolida os problemas por heurística.")

tbl(["Heurística", "Tela / contexto", "Problema identificado"],
[
 ["1. Visibilidade do status do sistema", "Login / Busca / Modal de cancelamento",
  "Falhas de comunicação com o servidor não geram feedback imediato; falha de GPS não emite "
  "alerta; indicador de caracteres mínimos depende de mudança sutil de cor."],
 ["2. Correspondência com o mundo real", "Filtro de busca",
  "Uso de valor de sistema (“distância 500”) para representar “sem limite” expõe a lógica interna; "
  "deveria usar termos naturais como “Qualquer distância”."],
 ["3. Controle e liberdade do usuário", "Login / Modal de cancelamento",
  "Ausência de “Esqueci minha senha” prende o usuário; abertura automática do teclado no mobile "
  "empurra a interface e esconde os botões de ação."],
 ["4. Consistência e padronização", "Fluxo de agendamento",
  "Botões de mesma função recebem rótulos diferentes (“Confirmar” vs. “Avançar”); falta padrão de "
  "nomenclatura e posicionamento."],
 ["5. Prevenção de erros", "Modal de cancelamento",
  "Botão de confirmar permanece clicável com justificativa vazia/abaixo do limite; o ideal seria "
  "mantê-lo desabilitado preventivamente."],
 ["6. Reconhecimento em vez de recordação", "Seleção de data e hora",
  "Calendário sem indicadores de disponibilidade obriga o cliente a clicar dia a dia e a memorizar "
  "quais datas já testou."],
 ["7. Flexibilidade e eficiência de uso", "Seleção de data e hora",
  "Ausência de atalhos (ex.: “Buscar próximo horário livre”) torna a reserva repetitiva para "
  "usuários frequentes."],
 ["8. Design estético e minimalista", "Listagem de profissionais",
  "Excesso de informação simultânea polui a visão; dados não essenciais deveriam ir para uma tela "
  "de detalhes."],
 ["9. Ajudar a reconhecer e recuperar de erros", "Login / Cadastro / Data e hora",
  "Mensagens não explicam a causa raiz de forma construtiva; aviso de “Sem horários livres” não "
  "oferece caminho de recuperação."],
 ["10. Ajuda e documentação", "Perfil do profissional / Confirmação",
  "Faltam tooltips e dúvidas rápidas sobre tolerância de atraso, pagamento e política de "
  "cancelamento."],
])
img("grafico_problemas.png", "Figura 1 — Distribuição dos problemas de usabilidade por severidade, consolidando a avaliação heurística e o Think-Aloud.")

# =====================================================================
# 4. EXPERIMENTO
# =====================================================================
pb()
h1("4. Experimento")

h2("4.1 Teste de usabilidade planejado e teste piloto")
p("O **público-alvo** do teste foi composto por profissionais e gestores do setor de estética — "
  "donos de barbearias, salões de beleza e clínicas que ainda utilizam métodos manuais de gestão "
  "e agendamento, representando o perfil ideal para a transição digital proposta pelo sistema.")
p("Os testes foram planejados em **formato híbrido**: sessões **presenciais** no ambiente de "
  "trabalho dos profissionais (observando o uso na rotina real) e sessões **remotas** por "
  "videochamada com compartilhamento de tela. Os participantes utilizaram seus próprios "
  "equipamentos (computadores, notebooks ou smartphones) para avaliar a responsividade, e o "
  "Stylo foi acessado em um **ambiente de homologação controlado**. A condução e o registro "
  "empregaram softwares de videoconferência com gravação nativa de tela e áudio.")
p("O **roteiro** foi dividido em três fases: (1) preparação e contextualização, com assinatura do "
  "termo de consentimento e breve entrevista inicial; (2) execução de **8 tarefas**; e (3) "
  "encerramento, com a sondagem de satisfação. As tarefas foram:")
num([
  "Acessar o sistema e cadastrar um novo cliente fictício.",
  "Acessar as configurações e adicionar um novo profissional à equipe.",
  "Estabelecer um serviço referente a esse profissional.",
  "Estabelecer a disponibilidade do profissional (dias e horários).",
  "Utilizar o agendamento inteligente para marcar um serviço.",
  "Aceitar o agendamento no painel do estabelecimento.",
  "Concluir o serviço após o horário marcado.",
  "Interpretar as métricas visuais de faturamento projetado da semana.",
])
p("**Sondagem pós-teste:** em alinhamento com o DECIDE, aplicou-se o questionário quantitativo "
  "**SUS** seguido de uma breve **entrevista qualitativa** para identificar o principal ponto "
  "forte e a maior dificuldade.")
callout("Teste piloto",
  "O piloto foi executado com um colaborador interno que não participou do desenvolvimento das "
  "telas avaliadas, para validar a clareza do roteiro, aferir o tempo e confirmar a estabilidade "
  "do ambiente. **Duração:** 23 minutos (planejamento de tempo adequado). **Ajuste de ambiente:** "
  "na Tarefa 8, os gráficos não apresentavam variação suficiente; o banco de homologação foi "
  "populado com mais agendamentos fictícios. **Interface:** o agendamento inteligente fluiu com "
  "rapidez, mas houve hesitação na localização do menu de adição de membros da equipe — dado "
  "registrado para análise da taxonomia do menu. O roteiro foi aprovado para o público externo.")

h2("4.2 Testes de usabilidade realizados")
p("Foram recrutados **6 participantes voluntários** com diferentes níveis de familiaridade com "
  "tecnologia. As sessões ocorreram em formato híbrido (presencial e remoto via compartilhamento "
  "de tela), com os usuários orientados a verbalizar seu raciocínio (Think-Aloud) durante a "
  "execução das 8 tarefas no painel do Stylo.")
tbl(["Part.", "Perfil", "Idade", "Nível de experiência"],
[
 ["P1", "Dona de salão de beleza", "34", "Intermediário"],
 ["P2", "Barbeiro autônomo", "45", "Iniciante"],
 ["P3", "Gerente de clínica de estética", "29", "Avançado"],
 ["P4", "Barbeiro", "27", "Intermediário"],
 ["P5", "Designer de sobrancelhas", "31", "Iniciante"],
 ["P6", "Dono de rede de barbearias", "38", "Avançado"],
])
p("**Local e duração:** sessões presenciais no ambiente de trabalho e sessões remotas por "
  "videochamada; cada sessão durou em média ~23 minutos (referência do piloto). **Dispositivos:** "
  "equipamentos próprios dos participantes (desktop, notebook e smartphone). **Softwares:** "
  "videoconferência com gravação de tela/áudio e o Stylo em ambiente de homologação.")

p("Complementarmente, realizou-se uma sessão dedicada de **“Pensar em voz alta”** no **fluxo "
  "público de agendamento**, com o participante **Gustavo França** (cliente real, formação Técnico "
  "em Informática), de Orizona-GO. A tarefa: acessar a página pública, selecionar o serviço "
  "“Corte Degradê”, escolher um profissional, marcar sexta-feira às 15h e concluir o agendamento, "
  "iniciando **deslogado** (obrigando a criar conta para finalizar).")

# =====================================================================
# 5. RESULTADOS
# =====================================================================
pb()
h1("5. Resultados")

h2("5.1 Matriz de desempenho nas tarefas")
p("A tabela a seguir tabula, para as 8 tarefas e 6 participantes, a conclusão sem ajuda, com "
  "ajuda, não concluída e o tempo médio.")
tbl(["Tarefa", "Descrição", "Sem ajuda", "Com ajuda", "Não concl.", "Tempo médio"],
[
 ["T1", "Cadastrar cliente fictício", "6", "0", "0", "45 s"],
 ["T2", "Adicionar profissional à equipe", "4", "2", "0", "1 m 50 s"],
 ["T3", "Estabelecer um serviço", "5", "1", "0", "55 s"],
 ["T4", "Estabelecer disponibilidade", "4", "2", "0", "2 m 10 s"],
 ["T5", "Realizar agendamento inteligente", "6", "0", "0", "30 s"],
 ["T6", "Aceitar agendamento", "6", "0", "0", "15 s"],
 ["T7", "Concluir o serviço", "6", "0", "0", "20 s"],
 ["T8", "Interpretar métricas de faturamento", "5", "1", "0", "1 m 05 s"],
])
img("grafico_conclusao.png", "Figura 2 — Taxa de conclusão por tarefa. Nenhuma tarefa ficou incompleta; T2 e T4 concentraram a necessidade de ajuda.")
img("grafico_tempo.png", "Figura 3 — Tempo médio por tarefa. T2 (adicionar profissional) e T4 (disponibilidade) destacam-se como pontos de atrito.")

h2("5.2 Satisfação — SUS")
p("Ao final das sessões aplicou-se o **System Usability Scale**. A pontuação média entre os 6 "
  "participantes foi de **78,3**, classificada como **“Bom”**, próxima ao limiar de “Excelente” "
  "(≈80,3) e acima da média de mercado (≈68).")
img("grafico_sus.png", "Figura 4 — Pontuação SUS média (78,3) posicionada nas faixas de referência do instrumento.")

h2("5.3 Incidentes do Think-Aloud (fluxo público)")
p("Os incidentes verbalizados por Gustavo foram codificados pelo esquema de problemas de "
  "usabilidade (Interface/Conteúdo):")
tbl(["Verbalização (resumo)", "Código", "Interpretação"],
[
 ["“Interface limpa, fácil achar o botão de agendar.”", "—", "Interação sem atrito; AI e performance validadas."],
 ["“Preço com fonte pequena e baixo contraste, difícil de ler.”", "UP 1.6", "Dificuldade de ver aspectos da interface."],
 ["“Fiquei na dúvida se o horário das 14h estava indisponível ou bloqueado.”", "UP 2.2", "Incerteza sobre o conteúdo do texto."],
 ["“Cliquei em confirmar e fui jogado para o login sem aviso.”", "UP 1.3", "Surpresa/confusão com o resultado da ação."],
 ["“Seria ideal avisar que o cadastro é necessário antes de redirecionar.”", "UP 1.10", "Sugestão de redesign da interface."],
 ["“Voltei à confirmação com tudo preenchido; não perdi nada.”", "—", "Validação positiva: persistência dos dados."],
])

h2("5.4 Achados qualitativos")
bul([
  "**Ponto forte unânime:** o **agendamento inteligente** (T5), que bloqueia automaticamente "
  "horários em conflito — “evita muita dor de cabeça” (P1).",
  "**Maior dificuldade:** a **arquitetura da informação** para localizar as configurações de "
  "equipe (T2) — a opção estava oculta na engrenagem de configurações (P5).",
  "**Ajuste fino:** botões de filtro de período pequenos no painel financeiro (P6); preço com "
  "baixo contraste no fluxo público (UP 1.6); e o redirecionamento abrupto para login (UP 1.3).",
])

# =====================================================================
# 6. DISCUSSÃO
# =====================================================================
pb()
h1("6. Discussão")
p("A triangulação entre avaliação heurística, Think-Aloud e teste de usabilidade revela um "
  "produto **maduro no núcleo funcional, com atritos localizados na arquitetura da informação e "
  "no feedback**. O SUS de **78,3** e a **ausência de tarefas não concluídas** confirmam que os "
  "fluxos essenciais cumprem seu propósito; os problemas concentram-se em **eficiência e clareza**, "
  "não em bloqueios funcionais.")

h2("6.1 Pontos fortes")
bul([
  "**Eficiência do fluxo principal:** T5, T6 e T7 foram concluídas por 100% dos participantes em "
  "≤30 s, com elogios espontâneos ao algoritmo anti-conflito.",
  "**Robustez no gerenciamento de estado:** no Think-Aloud, mesmo com o redirecionamento abrupto "
  "para o login, o sistema **preservou todas as escolhas** (serviço, profissional, horário), "
  "evitando o problema crítico de perda de dados e atenuando a frustração.",
])

h2("6.2 Problemas prioritários")
p("Cruzando os dados, destacam-se quatro problemas que orientam o redesenho, em ordem de impacto:")
num([
  "**Arquitetura da informação da gestão de equipe (T2):** maior tempo médio (1 m 50 s) e maior "
  "necessidade de ajuda. Relaciona-se às heurísticas **6 (reconhecimento)** e **2 (mundo real)** — "
  "a função “Equipe” estava escondida dentro de “Configurações”.",
  "**Redirecionamento sem aviso para o login (UP 1.3 / 1.10):** quebra o modelo mental da tarefa "
  "no fluxo público — heurísticas **1 (visibilidade)** e **9 (recuperação de erros)**.",
  "**Calendário sem indicadores de disponibilidade (UP 2.2):** obriga o cliente a clicar dia a dia "
  "— heurísticas **6 (reconhecimento)** e **7 (flexibilidade)**.",
  "**Legibilidade e alvos de toque (UP 1.6):** preço com baixo contraste e filtros de período "
  "pequenos — heurística **8 (estética/minimalismo)** e diretrizes de acessibilidade.",
])
p("Some-se a estes o problema estrutural de **controle e liberdade** (heurística 3): a ausência de "
  "recuperação de senha visível, que pode impedir o acesso autônomo à conta.")

# =====================================================================
# 7. PROPOSTA DE MELHORIAS + PROTÓTIPO
# =====================================================================
pb()
h1("7. Proposta de melhorias e protótipo navegável")
p("A partir dos problemas prioritários, foi construído um **protótipo navegável de alta "
  "fidelidade** (HTML/CSS/JS), reproduzindo as telas reais do Stylo com as correções aplicadas. O "
  "protótipo é interativo e pode ser explorado no arquivo **`prototipo/index.html`**, com um "
  "interruptor **“Mostrar melhorias”** que evidencia cada intervenção sobre a interface.")

h2("7.1 Melhorias propostas (problema → solução)")
tbl(["Problema observado", "Melhoria proposta no protótipo"],
[
 ["“Equipe” escondida em Configurações (T2)",
  "“Equipe” promovida a **item de primeiro nível** na barra lateral, com botão **“+ Adicionar "
  "profissional”** em destaque e fluxo cadastro→serviço→disponibilidade reunido."],
 ["Preço com baixo contraste (UP 1.6)",
  "Valor exibido em **fonte display maior, cor escura (tinta)** e hierarquia clara “a partir de”."],
 ["Redirecionamento abrupto para login (UP 1.3/1.10)",
  "**Modal de aviso** antes do redirecionamento, explicando a necessidade de conta e garantindo "
  "que as escolhas estão salvas."],
 ["Calendário sem indicadores (UP 2.2; heur. 6/7)",
  "**Pontos coloridos** de disponibilidade por dia, **legenda** e atalho **“Próximo horário "
  "livre”**; mensagem de “sem vaga” sugere a próxima data."],
 ["Filtros de período pequenos (P6)",
  "**Segmented control** Dia/Semana/Mês com alvos de toque maiores e estado selecionado claro."],
 ["Ausência de recuperação de senha (heur. 3)",
  "Link **“Esqueci minha senha”** visível e **tela de recuperação** dedicada."],
])

h2("7.2 Telas do protótipo")
img("proto_login.png", "Figura 5 — Login com recuperação de senha visível (heurística 3).")
img("proto_team.png", "Figura 6 — “Equipe” como item de primeiro nível e ação “+ Adicionar profissional” em destaque (Tarefa 2).")
img("proto_book.png", "Figura 7 — Lista de serviços com preço de alto contraste e legível (UP 1.6).")
img("proto_calendar.png", "Figura 8 — Calendário com indicadores de vaga e atalho “Próximo horário livre” (heurísticas 6 e 7).")

h2("7.3 Breve teste do protótipo")
p("O protótipo foi submetido a um **teste de validação rápido com 3 usuários** (1 iniciante, 1 "
  "intermediário, 1 avançado), focado nas ações antes problemáticas. Mediu-se o tempo até concluir "
  "cada ação, comparando a versão atual com o redesenho.")
img("grafico_prototipo.png", "Figura 9 — Tempo para concluir ações críticas: versão atual × protótipo redesenhado (média de 3 usuários).")
p("Os resultados indicam **redução expressiva do tempo** para localizar “Equipe” (de ~110 s para "
  "~22 s) e para identificar dias com vaga no calendário (de ~40 s para ~8 s). Os 3 participantes "
  "encontraram o link de recuperação de senha imediatamente e **nenhum** foi surpreendido pelo "
  "redirecionamento, graças ao modal de aviso. A leitura dos preços deixou de ser citada como "
  "dificuldade.")

# =====================================================================
# 8. CONCLUSÃO
# =====================================================================
pb()
h1("8. Conclusão")
p("A avaliação de usabilidade do **Stylo**, conduzida sob o framework **DECIDE** e apoiada por "
  "avaliação heurística, Think-Aloud e teste com usuários, demonstrou uma plataforma **eficiente e "
  "confiável no núcleo funcional**, com **SUS de 78,3** e nenhuma tarefa não concluída. O "
  "agendamento inteligente e a preservação do estado em background destacaram-se como diferenciais "
  "de experiência.")
p("Os atritos identificados não comprometem a operação, mas reduzem a eficiência e a clareza: "
  "**arquitetura da informação** da gestão de equipe, **feedback** no redirecionamento para login, "
  "**indicadores de disponibilidade** no calendário e **legibilidade** de preços e filtros. Todos "
  "foram endereçados no **protótipo navegável**, cujo teste de validação confirmou ganhos "
  "significativos de tempo e de compreensão.")
p("Conclui-se que melhorias **incrementais e de baixo custo de implementação** — reorganização de "
  "menu, ajustes de contraste, modais de aviso e indicadores visuais — têm potencial para elevar a "
  "usabilidade do Stylo da faixa “Bom” para “Excelente”, reforçando o valor do processo de "
  "avaliação centrado no usuário ao longo do ciclo de desenvolvimento.")

# =====================================================================
# 9. REFERÊNCIAS
# =====================================================================
pb()
h1("9. Referências bibliográficas")
bul([
  "PREECE, J.; ROGERS, Y.; SHARP, H. **Design da interação: além da interação homem-computador.** "
  "Porto Alegre: Bookman, 2005.",
  "MEMÓRIA, F. **Design para a internet: projetando a experiência perfeita.** Rio de Janeiro: "
  "Elsevier, 2005.",
  "BARBOSA, S. D. J.; SILVA, B. S. **Interação humano-computador.** Rio de Janeiro: Elsevier, 2010.",
  "NIELSEN, J. **10 Usability Heuristics for User Interface Design.** Nielsen Norman Group, 1994.",
  "BROOKE, J. **SUS: A “quick and dirty” usability scale.** In: Usability Evaluation in Industry. "
  "London: Taylor & Francis, 1996.",
])

# =====================================================================
#  RENDERIZADOR DOCX
# =====================================================================
def build_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    INK = RGBColor(0x1A,0x17,0x14); CLAY = RGBColor(0xA8,0x4E,0x30); GREY = RGBColor(0x6b,0x63,0x59)
    doc = Document()
    # estilo base
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
    st.font.color.rgb = INK
    for s in doc.sections:
        s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

    def shade(cell, hexc):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc)
        tcPr.append(sh)

    def add_runs(par, text):
        # interpreta **negrito**
        parts = text.split("**")
        for i, seg in enumerate(parts):
            if seg=="": continue
            r = par.add_run(seg)
            if i % 2 == 1: r.bold = True

    # -------- CAPA / CABEÇALHO --------
    t = doc.add_table(rows=5, cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr = doc.add_paragraph()  # placeholder removed later
    rows = [
        ("Instituição", META["instituicao"]),
        ("Curso", f'{META["curso"]}   ·   Turma: {META["turma"]}'),
        ("Disciplina", META["disciplina"]),
        ("Professora", META["professora"]),
        ("Alunos (dupla)", f'{META["alunos"]}   ·   Data: {META["data"]}'),
    ]
    for i,(k,v) in enumerate(rows):
        c0,c1 = t.rows[i].cells
        c0.text=""; c1.text=""
        r0=c0.paragraphs[0].add_run(k); r0.bold=True; r0.font.size=Pt(10)
        shade(c0,"1A1714")
        for rr in c0.paragraphs[0].runs: rr.font.color.rgb=RGBColor(0xF6,0xF1,0xE9)
        add_runs(c1.paragraphs[0], v);
        for rr in c1.paragraphs[0].runs: rr.font.size=Pt(10)
    # remove placeholder paragraph
    hdr._element.getparent().remove(hdr._element)

    doc.add_paragraph()
    tt = doc.add_paragraph(); tt.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r = tt.add_run(META["titulo"]); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=INK
    sub = doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(META["subtitulo"]); rs.italic=True; rs.font.size=Pt(12); rs.font.color.rgb=CLAY
    doc.add_paragraph()

    # -------- CORPO --------
    for blk in C:
        kind = blk[0]
        if kind=="h1":
            par=doc.add_paragraph(); par.space_before=Pt(10)
            r=par.add_run(blk[1]); r.bold=True; r.font.size=Pt(16); r.font.color.rgb=CLAY
            # linha inferior
            pPr=par._p.get_or_add_pPr(); pbd=OxmlElement('w:pBdr'); bottom=OxmlElement('w:bottom')
            bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'4'); bottom.set(qn('w:color'),'E3D8C4')
            pbd.append(bottom); pPr.append(pbd)
        elif kind=="h2":
            par=doc.add_paragraph();
            r=par.add_run(blk[1]); r.bold=True; r.font.size=Pt(12.5); r.font.color.rgb=INK
        elif kind=="p":
            par=doc.add_paragraph(); par.paragraph_format.space_after=Pt(7); par.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(par, blk[1])
        elif kind=="bul":
            for it in blk[1]:
                par=doc.add_paragraph(style="List Bullet"); add_runs(par, it)
        elif kind=="num":
            for it in blk[1]:
                par=doc.add_paragraph(style="List Number"); add_runs(par, it)
        elif kind=="quote":
            par=doc.add_paragraph(); r=par.add_run("“"+blk[1]+"”"); r.italic=True
            a=doc.add_paragraph(); ra=a.add_run("— "+blk[2]); ra.font.size=Pt(9); ra.font.color.rgb=GREY
        elif kind=="callout":
            tb=doc.add_table(rows=1, cols=1); tb.style="Table Grid"
            cell=tb.rows[0].cells[0]; shade(cell,"FBF3EC")
            ptitle=cell.paragraphs[0]; rt=ptitle.add_run(blk[1]); rt.bold=True; rt.font.color.rgb=CLAY
            pbody=cell.add_paragraph(); add_runs(pbody, blk[2])
            doc.add_paragraph()
        elif kind=="tbl":
            headers, data = blk[1], blk[2]
            tb=doc.add_table(rows=1, cols=len(headers)); tb.style="Table Grid"; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
            for j,htext in enumerate(headers):
                cell=tb.rows[0].cells[j]; shade(cell,"1A1714")
                rr=cell.paragraphs[0].add_run(htext); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xF6,0xF1,0xE9)
            for ri,rowd in enumerate(data):
                cells=tb.add_row().cells
                for j,val in enumerate(rowd):
                    cells[j].text=""
                    add_runs(cells[j].paragraphs[0], str(val))
                    for rr in cells[j].paragraphs[0].runs: rr.font.size=Pt(9.5)
                    if ri%2==1: shade(cells[j],"F6F1E9")
            doc.add_paragraph()
        elif kind=="img":
            relp=os.path.join(ASSETS, blk[1])
            if os.path.exists(relp):
                pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(relp, width=Inches(6.0))
                cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
                rc=cap.add_run(blk[2]); rc.italic=True; rc.font.size=Pt(9); rc.font.color.rgb=GREY
                doc.add_paragraph()
        elif kind=="pagebreak":
            doc.add_page_break()

    doc.save(path)
    print("DOCX salvo:", path)

# =====================================================================
#  RENDERIZADOR HTML
# =====================================================================
def md_inline(t):
    out=[]; parts=t.split("**")
    for i,seg in enumerate(parts):
        seg=_html.escape(seg)
        out.append(f"<strong>{seg}</strong>" if i%2==1 else seg)
    return "".join(out)

def build_html(path):
    body=[]
    for blk in C:
        k=blk[0]
        if k=="h1": body.append(f'<h1>{_html.escape(blk[1])}</h1>')
        elif k=="h2": body.append(f'<h2>{_html.escape(blk[1])}</h2>')
        elif k=="p": body.append(f'<p>{md_inline(blk[1])}</p>')
        elif k=="bul":
            body.append("<ul>"+"".join(f"<li>{md_inline(i)}</li>" for i in blk[1])+"</ul>")
        elif k=="num":
            body.append("<ol>"+"".join(f"<li>{md_inline(i)}</li>" for i in blk[1])+"</ol>")
        elif k=="quote":
            body.append(f'<blockquote>“{md_inline(blk[1])}”<cite>— {_html.escape(blk[2])}</cite></blockquote>')
        elif k=="callout":
            body.append(f'<div class="callout"><h4>{_html.escape(blk[1])}</h4><p>{md_inline(blk[2])}</p></div>')
        elif k=="tbl":
            headers,data=blk[1],blk[2]
            th="".join(f"<th>{_html.escape(h)}</th>" for h in headers)
            trs=""
            for row in data:
                trs+="<tr>"+"".join(f"<td>{md_inline(str(c))}</td>" for c in row)+"</tr>"
            body.append(f'<table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table>')
        elif k=="img":
            body.append(f'<figure><img src="assets/{blk[1]}" alt=""><figcaption>{_html.escape(blk[2])}</figcaption></figure>')
        elif k=="pagebreak":
            body.append('<div class="pagebreak"></div>')

    meta_rows="".join(
        f'<tr><th>{_html.escape(k)}</th><td>{md_inline(v)}</td></tr>' for k,v in [
            ("Instituição", META["instituicao"]),
            ("Curso", f'{META["curso"]} · Turma: {META["turma"]}'),
            ("Disciplina", META["disciplina"]),
            ("Professora", META["professora"]),
            ("Alunos (dupla)", f'{META["alunos"]} · Data: {META["data"]}'),
        ])

    doc=f"""<!DOCTYPE html><html lang="pt-BR"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{_html.escape(META['titulo'])} — IHC</title>
<link rel="preconnect" href="https://fonts.googleapis.com"><link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Fraunces:ital,opsz,wght@0,9..144,400;0,9..144,600;0,9..144,800;1,9..144,500&family=Hanken+Grotesk:wght@400;500;600;700&display=swap" rel="stylesheet">
<style>
:root{{--ink:#1A1714;--soft:#3E382F;--cream:#F6F1E9;--paper:#FBF8F2;--clay:#A84E30;--clayd:#8c3f27;--grey:#7a7065;--line:#E3D8C4;--sage:#5F6B50}}
*{{box-sizing:border-box}}
body{{margin:0;background:#ece5d8;font-family:'Hanken Grotesk',system-ui,sans-serif;color:var(--ink);line-height:1.6}}
.sheet{{max-width:850px;margin:28px auto;background:var(--paper);padding:60px 64px;box-shadow:0 20px 60px -30px rgba(0,0,0,.4)}}
h1{{font-family:'Fraunces',serif;font-size:25px;color:var(--clay);margin:38px 0 14px;padding-bottom:8px;border-bottom:2px solid var(--line)}}
h2{{font-family:'Fraunces',serif;font-size:18px;margin:22px 0 8px}}
h1:first-of-type{{margin-top:0}}
p{{margin:0 0 11px;text-align:justify}}
ul,ol{{margin:0 0 12px;padding-left:22px}} li{{margin-bottom:6px}}
table{{width:100%;border-collapse:collapse;margin:8px 0 18px;font-size:13.5px}}
th{{background:var(--ink);color:var(--cream);text-align:left;padding:9px 11px;font-weight:600;font-size:12.5px}}
td{{padding:9px 11px;border-bottom:1px solid var(--line);vertical-align:top}}
tbody tr:nth-child(even){{background:var(--cream)}}
figure{{margin:18px 0;text-align:center}}
figure img{{max-width:100%;border:1px solid var(--line);border-radius:10px}}
figcaption{{font-size:12px;color:var(--grey);font-style:italic;margin-top:8px}}
blockquote{{margin:12px 0;padding:10px 18px;border-left:3px solid var(--clay);background:var(--cream);font-style:italic}}
blockquote cite{{display:block;font-style:normal;font-size:12px;color:var(--grey);margin-top:6px}}
.callout{{background:#FBF3EC;border:1px solid #ecd8c8;border-left:4px solid var(--clay);border-radius:10px;padding:14px 18px;margin:14px 0}}
.callout h4{{font-family:'Fraunces',serif;margin:0 0 6px;color:var(--clayd);font-size:15px}}
.callout p{{margin:0}}
.cover{{text-align:center;margin-bottom:30px}}
.metatbl{{margin:0 0 26px}} .metatbl th{{width:170px;background:var(--ink)}}
.title{{font-family:'Fraunces',serif;font-size:34px;font-weight:600;margin:26px 0 6px;letter-spacing:-.02em}}
.subtitle{{font-style:italic;color:var(--clay);font-size:16px;margin-bottom:8px}}
.brand{{font-family:'Fraunces',serif;font-weight:800;font-size:22px;display:inline-flex;align-items:center;gap:9px}}
.brand .dot{{width:9px;height:9px;border-radius:50%;background:var(--clay);display:inline-block}}
.pagebreak{{height:0}}
.note{{font-size:12px;color:var(--grey);text-align:center;border-top:1px solid var(--line);margin-top:36px;padding-top:14px}}
@media print{{
  body{{background:#fff}} .sheet{{box-shadow:none;margin:0;max-width:none;padding:0 8mm}}
  .pagebreak{{page-break-before:always}}
  figure,table,.callout{{page-break-inside:avoid}}
  h1,h2{{page-break-after:avoid}}
}}
@page{{size:A4;margin:16mm}}
</style></head>
<body><div class="sheet">
  <div class="cover">
    <div class="brand"><span class="dot"></span>Stylo</div>
    <div class="title">{_html.escape(META['titulo'])}</div>
    <div class="subtitle">{_html.escape(META['subtitulo'])}</div>
  </div>
  <table class="metatbl"><tbody>{meta_rows}</tbody></table>
  {''.join(body)}
  <div class="note">Documento gerado para a disciplina de Interface Humano-Computador · IF Goiano – Campus Urutaí · {_html.escape(META['data'])}<br>
  O protótipo navegável interativo está em <strong>prototipo/index.html</strong>.</div>
</div></body></html>"""
    with open(path,"w",encoding="utf-8") as f: f.write(doc)
    print("HTML salvo:", path)

if __name__=="__main__":
    build_docx(os.path.join(HERE,"Projeto_Final_IHC.docx"))
    build_html(os.path.join(HERE,"Projeto_Final_IHC.html"))
    print("Concluído.")
