# -*- coding: utf-8 -*-
"""Gera a atividade 'Prototipação' em .docx (padrão IF Goiano)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
INK = RGBColor(0x1A,0x17,0x14); CLAY = RGBColor(0xA8,0x4E,0x30); GREY = RGBColor(0x6b,0x63,0x59)

doc = Document()
st = doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11); st.font.color.rgb=INK
for s in doc.sections:
    s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2); s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

def shade(cell,hexc):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexc); tcPr.append(sh)
def add_runs(par,text):
    for i,seg in enumerate(text.split("**")):
        if seg=="":continue
        r=par.add_run(seg)
        if i%2==1:r.bold=True
def H1(t):
    par=doc.add_paragraph(); par.space_before=Pt(10)
    r=par.add_run(t); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=CLAY
    pPr=par._p.get_or_add_pPr(); pbd=OxmlElement('w:pBdr'); b=OxmlElement('w:bottom')
    b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6'); b.set(qn('w:space'),'4'); b.set(qn('w:color'),'E3D8C4')
    pbd.append(b); pPr.append(pbd)
def H2(t):
    par=doc.add_paragraph(); r=par.add_run(t); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=INK
def P(t):
    par=doc.add_paragraph(); par.paragraph_format.space_after=Pt(7); par.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(par,t)
def BUL(items):
    for it in items:
        par=doc.add_paragraph(style="List Bullet"); add_runs(par,it)
def NUM(items):
    for it in items:
        par=doc.add_paragraph(style="List Number"); add_runs(par,it)
def TBL(headers,rows):
    tb=doc.add_table(rows=1,cols=len(headers)); tb.style="Table Grid"; tb.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        c=tb.rows[0].cells[j]; shade(c,"1A1714")
        rr=c.paragraphs[0].add_run(h); rr.bold=True; rr.font.size=Pt(9.5); rr.font.color.rgb=RGBColor(0xF6,0xF1,0xE9)
    for ri,row in enumerate(rows):
        cells=tb.add_row().cells
        for j,v in enumerate(row):
            cells[j].text=""; add_runs(cells[j].paragraphs[0],str(v))
            for rr in cells[j].paragraphs[0].runs: rr.font.size=Pt(9.5)
            if ri%2==1: shade(cells[j],"F6F1E9")
    doc.add_paragraph()

# ---------------- CABEÇALHO ----------------
meta=[("Instituição","INSTITUTO FEDERAL GOIANO – CAMPUS URUTAÍ"),
      ("Curso","Sistemas de Informação   ·   Turma: 7º Período"),
      ("Disciplina","Interface Humano-Computador (IHC)"),
      ("Professora","Cristiane de Fátima dos Santos Cardoso"),
      ("Alunos (dupla)","Arthur Faria Estrela e Sávio Issa de Sousa   ·   Data: 23/06/2026")]
t=doc.add_table(rows=len(meta),cols=2); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
for i,(k,v) in enumerate(meta):
    c0,c1=t.rows[i].cells; c0.text="";c1.text=""
    r0=c0.paragraphs[0].add_run(k); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=RGBColor(0xF6,0xF1,0xE9)
    shade(c0,"1A1714"); add_runs(c1.paragraphs[0],v)
    for rr in c1.paragraphs[0].runs: rr.font.size=Pt(10)
doc.add_paragraph()
tt=doc.add_paragraph(); tt.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tt.add_run("Atividade — Prototipação"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=INK
doc.add_paragraph()

# ============ 1 ============
H1("1. Design centrado no usuário")
P("O **Design Centrado no Usuário (DCU)** — em inglês *User-Centered Design* — é uma filosofia e um "
  "processo de projeto que coloca o **usuário no centro de todas as decisões de design**, do início ao "
  "fim do desenvolvimento. Em vez de partir das preferências do programador ou da tecnologia "
  "disponível, parte-se das **necessidades, objetivos, características e contexto reais das pessoas** "
  "que vão usar o sistema. Segundo Preece, Rogers e Sharp (2005), o objetivo é desenvolver produtos "
  "**usáveis e úteis**, que apoiem as atividades das pessoas com eficiência, eficácia e satisfação.")
H2("Princípios fundamentais")
BUL([
  "**Foco desde cedo nos usuários e nas tarefas:** entender quem são os usuários, o que precisam fazer "
  "e em que contexto, antes de começar a projetar.",
  "**Medição empírica:** observar e medir o desempenho e as reações dos usuários com protótipos e "
  "versões do produto, usando dados reais (e não suposições).",
  "**Design iterativo:** projetar, testar, ajustar e repetir; os problemas encontrados realimentam um "
  "novo ciclo de melhoria.",
])
P("Esse processo é formalizado pela norma **ISO 9241-210**, que descreve um ciclo iterativo com quatro "
  "atividades: (1) **entender e especificar o contexto de uso**; (2) **especificar os requisitos** do "
  "usuário; (3) **produzir soluções de design** (esboços e protótipos); e (4) **avaliar o design com "
  "usuários reais**. O ciclo se repete até que os requisitos de usabilidade sejam atingidos.")
P("**Benefícios:** produtos mais fáceis de aprender e de usar, maior satisfação, menos erros, menor "
  "necessidade de suporte e **redução de retrabalho** — pois os problemas são descobertos cedo, quando "
  "ainda são baratos de corrigir. No nosso projeto, o **Stylo** seguiu essa lógica: avaliamos a "
  "interface com usuários reais do setor de estética (Think-Aloud, teste de usabilidade e SUS) e "
  "usamos os problemas encontrados para guiar o redesenho — o usuário, e não o nosso gosto, definiu "
  "as melhorias.")

# ============ 2 ============
H1("2. Nível de fidelidade de um protótipo")
P("A **fidelidade** de um protótipo é o **grau de semelhança entre o protótipo e o produto final**, "
  "considerando aspectos visuais (aparência), de conteúdo e de interatividade. Escolher o nível certo "
  "depende da **fase do projeto** e do que se deseja aprender ou comunicar. Costuma-se falar em três "
  "níveis:")
TBL(["Nível","Como é","Quando usar","Vantagens / cuidados"],
[
 ["Baixa fidelidade","Esboços em papel, wireframes, rascunhos em quadro. Sem cores ou conteúdo real.",
  "Início do projeto, para explorar ideias e fluxos rapidamente.",
  "Rápido e barato; estimula críticas e mudanças. Não representa o visual final."],
 ["Média fidelidade","Wireframes digitais navegáveis, com estrutura e hierarquia, mas visual neutro.",
  "Para validar arquitetura da informação e navegação.",
  "Equilíbrio entre custo e realismo; ainda abstrai o estilo visual."],
 ["Alta fidelidade","Parece e funciona como o produto real: cores, tipografia, conteúdo e interações reais.",
  "Para testes de usabilidade detalhados, validação e apresentação a clientes.",
  "Resultados de teste mais precisos; porém é mais caro e demorado, e o usuário pode focar no visual."],
])
P("Além do eixo de fidelidade, os protótipos também variam quanto à **abrangência**: um protótipo "
  "**horizontal** mostra muitas funções superficialmente (visão geral da interface), enquanto um "
  "protótipo **vertical** detalha poucas funções em profundidade (um fluxo completo). Quanto ao "
  "**destino**, podem ser **descartáveis** (*throwaway*, feitos só para aprender e depois jogados "
  "fora) ou **evolutivos** (*evolutionary*, que evoluem até virar o produto).")
P("**Trade-off central:** quanto maior a fidelidade, mais realista o teste, porém maior o custo e o "
  "tempo — e maior o risco de o usuário (ou o cliente) achar que o produto “já está pronto” e resistir "
  "a mudanças. Por isso recomenda-se **começar em baixa fidelidade** e aumentar gradualmente. No "
  "**Stylo**, por exemplo, a reorganização do menu de “Equipe” poderia ser explorada primeiro em "
  "esboços de baixa fidelidade; já a validação final foi feita com um **protótipo navegável de alta "
  "fidelidade**, que reproduz as telas reais e a interação.")

# ============ 3 ============
H1("3. Passo a passo para criar um protótipo")
P("Antes de desenhar qualquer tela, vale considerar **pontos iniciais** que orientam todas as "
  "decisões seguintes:")
BUL([
  "**Qual é o objetivo do protótipo?** Explorar uma ideia, comunicar uma solução ou validar/testar uma hipótese?",
  "**Quem é o usuário?** Perfil, nível de experiência, necessidades e dores (personas).",
  "**Quais são as tarefas e os fluxos críticos** que o protótipo precisa cobrir?",
  "**O que queremos aprender?** Hipóteses claras tornam o teste objetivo.",
  "**Qual nível de fidelidade é adequado** à fase atual do projeto?",
  "**Quais são as restrições** de tempo, orçamento, ferramentas e limitações técnicas/de negócio?",
])
P("Definidos esses pontos, um **passo a passo** prático é:")
NUM([
  "**Entender o problema e definir objetivos:** o que o produto resolve e o que o protótipo precisa demonstrar ou validar.",
  "**Conhecer os usuários e o contexto:** levantar personas, cenários de uso e tarefas reais (pesquisa, entrevistas, observação).",
  "**Especificar requisitos e priorizar:** listar o que é essencial e selecionar os fluxos mais importantes para prototipar primeiro.",
  "**Escolher o nível de fidelidade e a ferramenta:** papel, Figma, Balsamiq, HTML, etc., conforme o objetivo e os recursos.",
  "**Esboçar a estrutura e a navegação:** definir a arquitetura da informação, os fluxos e o esqueleto das telas (wireframes).",
  "**Construir o protótipo:** montar as telas e ligar as interações, mantendo consistência visual e de comportamento.",
  "**Testar com usuários reais:** aplicar tarefas, observar, ouvir (Think-Aloud) e coletar dados de desempenho e satisfação.",
  "**Analisar e iterar:** corrigir os problemas encontrados e repetir o ciclo até atingir as metas de usabilidade.",
  "**Documentar e repassar (handoff):** registrar decisões, especificações e fluxos para a equipe de desenvolvimento.",
])
P("Em resumo, prototipar é um processo **iterativo e centrado no usuário**: começa-se simples e barato "
  "para aprender rápido, e aumenta-se a fidelidade à medida que as decisões se consolidam — exatamente "
  "o caminho que seguimos na avaliação e no redesenho do **Stylo**.")

# Referências
H1("Referências")
BUL([
  "PREECE, J.; ROGERS, Y.; SHARP, H. **Design da interação: além da interação homem-computador.** Porto Alegre: Bookman, 2005.",
  "BARBOSA, S. D. J.; SILVA, B. S. **Interação humano-computador.** Rio de Janeiro: Elsevier, 2010.",
  "MEMÓRIA, F. **Design para a internet: projetando a experiência perfeita.** Rio de Janeiro: Elsevier, 2005.",
  "ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. **ISO 9241-210 — Ergonomia da interação humano-sistema: Design centrado no humano para sistemas interativos.**",
])

out=os.path.join(HERE,"Atividade_Prototipacao.docx")
doc.save(out); print("salvo:",out)
